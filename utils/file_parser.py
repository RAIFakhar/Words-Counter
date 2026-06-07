"""
utils/file_parser.py
───────────────────────────────────────────────────────────────
FileParser — reads plain text from uploaded documents.

Supported formats:
  .txt   Plain text (UTF-8)
  .md    Markdown  (strips syntax before analysis)
  .docx  Word document (via python-docx)
  .pdf   PDF (via pdfplumber)

Usage:
    text, error = FileParser.parse(file_object, "document.pdf")
    if error:
        show_error(error)
    else:
        process(text)
"""
import re
from pathlib import Path


class FileParser:
    """Parse various document formats into plain text strings."""

    SUPPORTED = {".txt", ".md", ".docx", ".pdf"}

    @classmethod
    def parse(cls, file_obj, file_name: str) -> tuple[str, str]:
        """
        Parse ``file_obj`` (file-like object) given ``file_name``.

        Returns:
            (text_content, error_message)
            error_message is an empty string on success.
        """
        suffix = Path(file_name).suffix.lower()

        if suffix not in cls.SUPPORTED:
            return "", (
                f"Unsupported format '{suffix}'. "
                f"Accepted: {', '.join(sorted(cls.SUPPORTED))}"
            )

        try:
            dispatch = {
                ".txt":  cls._txt,
                ".md":   cls._md,
                ".docx": cls._docx,
                ".pdf":  cls._pdf,
            }
            return dispatch[suffix](file_obj), ""

        except ImportError as exc:
            pkg = str(exc).split("'")[-2] if "'" in str(exc) else str(exc)
            return "", (
                f"Missing library for '{suffix}' files: install '{pkg}'\n"
                f"  pip install {pkg}"
            )
        except Exception as exc:  # noqa: BLE001
            return "", f"Could not read '{file_name}': {exc}"

    # ─── Format handlers ──────────────────────────────────────────────────────

    @staticmethod
    def _txt(file_obj) -> str:
        raw = file_obj.read()
        return raw.decode("utf-8", errors="replace") if isinstance(raw, bytes) else raw

    @staticmethod
    def _md(file_obj) -> str:
        raw = file_obj.read()
        text = raw.decode("utf-8", errors="replace") if isinstance(raw, bytes) else raw

        # Strip common Markdown syntax so word-counts are meaningful
        text = re.sub(r"#{1,6}\s+",                       "",  text)          # headings
        text = re.sub(r"\*{1,3}([^*\n]+)\*{1,3}",        r"\1", text)        # bold/italic
        text = re.sub(r"_{1,2}([^_\n]+)_{1,2}",          r"\1", text)        # underscore italic
        text = re.sub(r"\[([^\]]+)\]\([^\)]*\)",          r"\1", text)        # links
        text = re.sub(r"!\[([^\]]*)\]\([^\)]*\)",         "",  text)          # images
        text = re.sub(r"`{1,3}[^`]*`{1,3}",               "",  text)          # inline / fenced code
        text = re.sub(r"```[\s\S]*?```",                   "",  text)          # fenced blocks
        text = re.sub(r"^[-*+]\s+",                        "",  text, flags=re.M)  # unordered list
        text = re.sub(r"^\d+\.\s+",                        "",  text, flags=re.M)  # ordered list
        text = re.sub(r"^>\s+",                            "",  text, flags=re.M)  # blockquotes
        text = re.sub(r"^---+$",                           "",  text, flags=re.M)  # horizontal rules
        text = re.sub(r"\|[^\n]+\|",                       "",  text)          # tables
        return text.strip()

    @staticmethod
    def _docx(file_obj) -> str:
        from docx import Document  # python-docx

        doc   = Document(file_obj)
        parts = []

        for para in doc.paragraphs:
            stripped = para.text.strip()
            if stripped:
                parts.append(stripped)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in parts:
                        parts.append(cell_text)

        return "\n\n".join(parts)

    @staticmethod
    def _pdf(file_obj) -> str:
        import pdfplumber  # pdfplumber

        pages = []
        with pdfplumber.open(file_obj) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    pages.append(page_text.strip())

        return "\n\n".join(pages)

    # ─── Utility ─────────────────────────────────────────────────────────────

    @staticmethod
    def preview(text: str, max_chars: int = 300) -> str:
        """Return a truncated preview of ``text``."""
        if len(text) <= max_chars:
            return text
        return text[:max_chars].rstrip() + "…"
